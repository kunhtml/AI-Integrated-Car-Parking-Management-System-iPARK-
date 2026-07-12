# -*- coding: utf-8 -*-
"""Generate UML-standard Use Case Specification Word document for iPARK."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "USE_CASE_SPECIFICATION_UML.docx"

from uc_data import USE_CASES  # noqa: E402

def classify_uc_complexity(main_transactions: int) -> str:
    """Phân loại theo Use Case Points (UCP): Simple ≤3, Medium 4–7, Complex >7."""
    if main_transactions <= 3:
        return "Simple"
    if main_transactions <= 7:
        return "Medium"
    return "Complex"


def get_transaction_stats(uc: dict) -> dict:
    main = len(uc["basic_flow"])
    alt = len(uc["alternative_flows"])
    exc = len(uc["exception_flows"])
    total = main + alt + exc
    complexity = classify_uc_complexity(main)
    return {
        "main": main,
        "alt": alt,
        "exc": exc,
        "total": total,
        "complexity": complexity,
    }


def enrich_use_cases() -> list[dict]:
    enriched = []
    for uc in USE_CASES:
        stats = get_transaction_stats(uc)
        enriched.append({**uc, **stats})
    return enriched


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_field_table(doc: Document, fields: list[tuple[str, str]]):
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12)

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        set_cell_shading(row.cells[0], "D9E2F3")
        p0 = row.cells[0].paragraphs[0]
        run0 = p0.add_run(label)
        run0.bold = True
        run0.font.size = Pt(10)
        run0.font.name = "Times New Roman"
        run0._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        p1 = row.cells[1].paragraphs[0]
        run1 = p1.add_run(value)
        run1.font.size = Pt(10)
        run1.font.name = "Times New Roman"
        run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return table


def add_numbered_section(doc: Document, title: str, items: list[str], empty_msg: str = "(Không có)"):
    doc.add_paragraph(title, style="Heading 3")
    if not items:
        p = doc.add_paragraph(empty_msg)
        p.paragraph_format.left_indent = Cm(0.5)
        return
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_document() -> Document:
    doc = Document()
    cases = enrich_use_cases()
    simple_count = sum(1 for uc in cases if uc["complexity"] == "Simple")
    medium_count = sum(1 for uc in cases if uc["complexity"] == "Medium")
    complex_count = sum(1 for uc in cases if uc["complexity"] == "Complex")

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ĐẶC TẢ USE CASE (UML)\n")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Hệ thống quản lý bãi đỗ xe thông minh iPARK\n\n")
    rs.font.size = Pt(14)
    rs.font.name = "Times New Roman"
    rs._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"Phiên bản: 3.0 (27 Medium UC)\n"
        f"Ngày: {date.today().strftime('%d/%m/%Y')}\n"
        f"Số lượng Use Case: {len(cases)} (UC01–UC27)\n"
        f"Medium UC: {medium_count} | Simple UC: {simple_count} | Complex UC: {complex_count}\n"
        f"Nguyên tắc: Vận hành TỰ ĐỘNG — kê khai xe tại UC06; mua vé tháng UC07 chỉ chọn biển\n"
        f"Mẫu: UML (Cockburn / RUP) + UCP + «include»\n"
    )
    rm.font.size = Pt(11)
    rm.font.name = "Times New Roman"
    rm._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_page_break()

    doc.add_heading("NGUYÊN TẮC VẬN HÀNH HỆ THỐNG", level=1)
    policy = doc.add_paragraph(
        "Hệ thống iPARK vận hành theo mô hình TỰ ĐỘNG HÓA TOÀN DIỆN:\n\n"
        "• Vào/ra bãi: Camera AI tự nhận diện biển số, tự cấp chỗ, tự tạo phiên, tự mở barrier.\n"
        "• Thanh toán: Tự tính phí, tự sinh VietQR, tự đồng bộ kết quả thanh toán.\n"
        "• Thành viên: Tự kích hoạt gói sau thanh toán; tự áp dụng ưu đãi khi vào/ra.\n"
        "• Giám sát: Tự phát hiện sự cố thiết bị/AI và tạo incident.\n\n"
        "THAO TÁC KÊ KHAI XE (phía khách hàng):\n"
        "Khách kê khai biển số và thông tin xe một lần tại UC06 (Quản lý phương tiện). "
        "Mua vé tháng (UC07) chỉ chọn biển đã kê khai — không nhập xe tại bước mua gói. "
        "Mọi bước vào/ra bãi đều TỰ ĐỘNG.\n\n"
        "Nhân viên (Staff) chỉ can thiệp khi hệ thống không tự xử lý được (UC14, UC16).\n"
        "Admin cấu hình một lần (UC17–UC27); vận hành hàng ngày tự động.\n\n"
        "Phân nhóm Actor:\n"
        "🅰️ Guest UC01–UC04 | 🅱️ Customer UC05–UC06, UC09–UC12 | Member UC07–UC08\n"
        "🅲 Staff UC13–UC16 | 🅳 Admin UC17–UC27"
    )
    for run in policy.runs:
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_page_break()

    # Table of contents heading
    doc.add_heading("MỤC LỤC USE CASE", level=1)
    toc_table = doc.add_table(rows=len(cases) + 1, cols=9)
    toc_table.style = "Table Grid"
    headers = [
        "Mã UC",
        "Nhóm Actor",
        "Tên Use Case",
        "UCP",
        "Main",
        "Alt",
        "Exc",
        "Tổng",
        "Tự động",
    ]
    for j, h in enumerate(headers):
        set_cell_shading(toc_table.rows[0].cells[j], "BDD7EE")
        toc_table.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    for i, uc in enumerate(cases, 1):
        auto_short = "Tự động" if "THỦ CÔNG" not in uc.get("automation", "") else "Bán tự động"
        toc_table.rows[i].cells[0].text = uc["id"]
        toc_table.rows[i].cells[1].text = uc.get("actor_group", "")
        toc_table.rows[i].cells[2].text = uc["name"]
        toc_table.rows[i].cells[3].text = uc["complexity"]
        toc_table.rows[i].cells[4].text = str(uc["main"])
        toc_table.rows[i].cells[5].text = str(uc["alt"])
        toc_table.rows[i].cells[6].text = str(uc["exc"])
        toc_table.rows[i].cells[7].text = str(uc["total"])
        toc_table.rows[i].cells[8].text = auto_short
        if uc["complexity"] == "Medium":
            set_cell_shading(toc_table.rows[i].cells[3], "FFF2CC")

    doc.add_page_break()

    # UCP summary section
    doc.add_heading("THỐNG KÊ PHÂN LOẠI UC & TRANSACTION (UCP)", level=1)
    summary = doc.add_paragraph(
        "Transaction = một bước tương tác giữa Actor và Hệ thống trong kịch bản Use Case.\n"
        "Phân loại độ phức tạp (theo số transaction luồng chính — Main):\n"
        "• Simple UC: 1–3 transaction\n"
        "• Medium UC: 4–7 transaction\n"
        "• Complex UC: ≥8 transaction\n\n"
        f"Tổng số Use Case: {len(cases)}\n"
        f"• Medium UC: {medium_count} use case\n"
        f"• Simple UC: {simple_count} use case\n"
        f"• Complex UC: {complex_count} use case"
    )
    for run in summary.runs:
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_paragraph()
    stats_table = doc.add_table(rows=len(cases) + 1, cols=10)
    stats_table.style = "Table Grid"
    stat_headers = [
        "Mã UC",
        "Nhóm",
        "Tên Use Case",
        "UCP",
        "Main",
        "Alt",
        "Exc",
        "Tổng",
        "Medium?",
        "Tự động",
    ]
    for j, h in enumerate(stat_headers):
        set_cell_shading(stats_table.rows[0].cells[j], "D9E2F3")
        stats_table.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    for i, uc in enumerate(cases, 1):
        note = "Medium UC" if uc["complexity"] == "Medium" else f"{uc['complexity']} UC"
        auto_short = "Tự động" if "THỦ CÔNG" not in uc.get("automation", "") else "Bán tự động"
        stats_table.rows[i].cells[0].text = uc["id"]
        stats_table.rows[i].cells[1].text = uc.get("actor_group", "")
        stats_table.rows[i].cells[2].text = uc["name"]
        stats_table.rows[i].cells[3].text = uc["complexity"]
        stats_table.rows[i].cells[4].text = str(uc["main"])
        stats_table.rows[i].cells[5].text = str(uc["alt"])
        stats_table.rows[i].cells[6].text = str(uc["exc"])
        stats_table.rows[i].cells[7].text = str(uc["total"])
        stats_table.rows[i].cells[8].text = note
        stats_table.rows[i].cells[9].text = auto_short
        if uc["complexity"] == "Medium":
            for col in range(10):
                set_cell_shading(stats_table.rows[i].cells[col], "FFF9E6")

    doc.add_page_break()

    # Template note
    doc.add_heading("HƯỚNG DẪN ĐỌC TÀI LIỆU", level=1)
    guide = (
        "Tài liệu tuân theo mẫu đặc tả Use Case UML chuẩn (Cockburn / RUP).\n\n"
        "Mỗi UC có: Nhóm Actor, Chế độ tự động, «include» sub-flows, Phân loại UCP, "
        "số Transaction (Main/Alt/Exc/Tổng), luồng chính/thay thế/ngoại lệ.\n\n"
        "«include»: use case con bắt buộc được gọi trong use case cha (UML include).\n\n"
        "Ký hiệu luồng thay thế/ngoại lệ: «bước»a. — ví dụ 3a. là nhánh tại bước 3."
    )
    p = doc.add_paragraph(guide)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_page_break()

    # Each use case
    for uc in cases:
        doc.add_heading(f"{uc['id']}: {uc['name']}", level=1)

        complexity_label = {
            "Simple": "Simple UC (1–3 transaction luồng chính)",
            "Medium": "Medium UC (4–7 transaction luồng chính)",
            "Complex": "Complex UC (≥8 transaction luồng chính)",
        }[uc["complexity"]]

        includes_text = (
            "\n".join(f"• {item}" for item in uc.get("includes", []))
            if uc.get("includes")
            else "(Không có «include»)"
        )

        fields = [
            ("Mã Use Case", uc["id"]),
            ("Tên Use Case", uc["name"]),
            ("Nhóm Actor", uc.get("actor_group", "")),
            ("Chế độ vận hành", uc.get("automation", "")),
            ("«include» Sub-flows", includes_text),
            ("Phân loại UC\n(Use Case Complexity)", complexity_label),
            (
                "Số Transaction\n(Transactions)",
                f"Luồng chính (Main): {uc['main']} | "
                f"Luồng thay thế (Alt): {uc['alt']} | "
                f"Luồng ngoại lệ (Exc): {uc['exc']} | "
                f"Tổng cộng: {uc['total']}",
            ),
            ("Tác nhân chính\n(Primary Actor)", uc["primary_actor"]),
            ("Tác nhân phụ\n(Secondary Actors)", uc["secondary_actors"]),
            ("Mô tả tóm tắt\n(Brief Description)", uc["summary"]),
            ("Điều kiện tiên quyết\n(Preconditions)", uc["preconditions"]),
            ("Điều kiện hậu\n(Postconditions)", uc["postconditions"]),
            ("Kích hoạt\n(Trigger)", uc["trigger"]),
            ("Yêu cầu đặc biệt\n(Special Requirements)", uc["special_requirements"]),
            ("Use Case liên quan\n(Related Use Cases)", uc["related_uc"]),
            ("Triển khai iPARK", uc["implementation"]),
        ]
        add_field_table(doc, fields)
        doc.add_paragraph()

        txn_note = doc.add_paragraph(
            f"Chi tiết transaction: {uc['main']} bước luồng chính + "
            f"{uc['alt']} nhánh thay thế + {uc['exc']} nhánh ngoại lệ = "
            f"{uc['total']} transaction tổng. "
            f"Đây là {complexity_label}."
        )
        for run in txn_note.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        if uc.get("includes"):
            add_numbered_section(
                doc,
                "«include» Sub-use Cases",
                [f"«include» {item}" for item in uc["includes"]],
            )

        doc.add_paragraph()
        add_numbered_section(
            doc,
            f"Luồng chính (Main Success Scenario) — {uc['main']} transaction",
            uc["basic_flow"],
        )
        add_numbered_section(
            doc,
            f"Luồng thay thế (Alternative Flows) — {uc['alt']} transaction",
            uc["alternative_flows"],
        )
        add_numbered_section(
            doc,
            f"Luồng ngoại lệ (Exception Flows) — {uc['exc']} transaction",
            uc["exception_flows"],
        )

        doc.add_paragraph()
        doc.add_page_break()

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
